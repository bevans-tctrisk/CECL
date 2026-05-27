"""Generate CECL Data Security & Privacy Documentation (Word format)."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h

def add_table_row(table, cells, bold=False, header=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(val)
        for p in cell.paragraphs:
            p.style.font.size = Pt(10)
            for run in p.runs:
                run.font.size = Pt(10)
                if bold or header:
                    run.font.bold = True
                if header:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if header:
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): '1B3A5C',
                qn('w:val'): 'clear',
            })
            shading.append(shd)
    return row

def build_document():
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CECL Credit Migration Analysis\nData Security & Privacy Policy")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Prepared for Regulatory Review")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Table of Contents placeholder ──
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. System Architecture Overview",
        "3. Data Classification",
        "4. Data Lifecycle Management",
        "5. Secure Data Transfer via Egnyte",
        "6. Access Controls & Authentication",
        "7. Data Processing Safeguards",
        "8. Output Report Controls",
        "9. External Data Transmission",
        "10. Data Retention & Disposal",
        "11. Security Controls & Recommendations",
        "Appendix A: Data Field Inventory",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document describes the data security controls and privacy safeguards "
        "implemented within the CECL (Current Expected Credit Losses) Credit Migration "
        "Analysis system. The system processes credit union member loan data to produce "
        "regulatory compliance reports, including credit risk migration matrices, "
        "historical balance trends, charge-off and recovery analysis, and CECL reserve "
        "calculations."
    )
    doc.add_paragraph(
        "The system handles limited categories of personally identifiable information (PII) "
        "and sensitive financial data. This document details how that data is ingested, "
        "processed, stored, and safeguarded throughout its lifecycle. It also identifies "
        "areas where additional controls are recommended."
    )

    add_heading_styled(doc, "Key Findings", level=2)
    bullets = [
        "All data processing occurs on a local, on-premises workstation — no member data "
        "is transmitted to third-party analytics platforms or external processing services.",
        "Source data is transferred securely from credit unions via Egnyte, an enterprise "
        "file management platform with SOC 2 Type II certification, AES-256 encryption "
        "at rest, and TLS/SSL encryption in transit (see Section 5).",
        "The system collects only the minimum data fields required for CECL compliance "
        "calculations: member identifiers, loan balances, and credit scores.",
        "No names, Social Security numbers, dates of birth, addresses, or contact "
        "information are collected, stored, or processed at any point.",
        "External API calls are limited to publicly available economic indicators "
        "(Bureau of Labor Statistics, U.S. Census Bureau) and transmit zero member data.",
        "Output reports contain aggregated pool-level data. Member-level detail is "
        "limited to a single worksheet used for audit verification purposes.",
        "The database resides on localhost and is not exposed to any network interface.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. System Architecture Overview", level=1)
    doc.add_paragraph(
        "The CECL Credit Migration Analysis system is a self-contained, locally-hosted "
        "application running on a single secured workstation. It consists of the following "
        "components:"
    )

    add_heading_styled(doc, "2.1 Component Overview", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Purpose"
    hdr[2].text = "PII Exposure"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)

    components = [
        ("Import Module\n(import_data.py)", "Reads source Excel/CSV files, normalizes data, "
         "loads into local database", "Processes member numbers\nand FICO scores"),
        ("CECL Engine\n(cecl_engine.py)", "Calculates risk migration matrices, reserve rates, "
         "and expected loss amounts", "Accesses member-level\nbalances and grades"),
        ("Report Generator\n(generate_report.py)", "Orchestrates data loading, calculations, "
         "and report creation", "Aggregates member data\ninto pool-level summaries"),
        ("TCT Report Builder\n(report_tct.py)", "Produces TCT-format Excel workbook with "
         "multiple analysis tabs", "Outputs aggregated data;\none audit detail sheet"),
        ("Vizo Report Builder\n(report_vizo.py)", "Produces Vizo-format Excel workbooks "
         "(main + supplemental)", "Outputs aggregated data;\none audit detail sheet"),
        ("Economic Data Fetcher\n(fetch_econ_data.py)", "Retrieves public economic indicators "
         "from federal APIs", "Zero PII — transmits only\ngeographic identifiers"),
        ("Local Database\n(PostgreSQL)", "Stores processed loan snapshots for historical "
         "trend analysis", "Stores member numbers\nand FICO scores locally"),
    ]
    for comp in components:
        add_table_row(table, comp)

    doc.add_paragraph()
    add_heading_styled(doc, "2.2 Data Flow Diagram", level=2)
    doc.add_paragraph(
        "The data flow is strictly linear and contained within the local workstation:"
    )
    flow_steps = [
        "Source Files (Excel/CSV) → uploaded by credit union to Egnyte (encrypted in transit "
        "via TLS/SSL, encrypted at rest via AES-256)",
        "Analyst Download → files retrieved from Egnyte to secured local workstation",
        "Import Module → parses, validates, and normalizes loan records",
        "Local PostgreSQL Database → stores processed snapshots (localhost only)",
        "Report Generator → queries database, calculates CECL metrics",
        "Output Reports (Excel) → written to secured network share or uploaded to Egnyte "
        "for secure delivery to credit union",
    ]
    for i, step in enumerate(flow_steps, 1):
        p = doc.add_paragraph(f"Step {i}: {step}")
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("File transfers between credit unions and the analysis environment are "
                     "conducted exclusively through Egnyte's encrypted platform. At no point "
                     "during processing does member data leave the local workstation or traverse "
                     "any unencrypted network.")
    run.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 3. DATA CLASSIFICATION
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Data Classification", level=1)
    doc.add_paragraph(
        "All data processed by the system is classified according to the following "
        "sensitivity tiers:"
    )

    add_heading_styled(doc, "3.1 Data Collected", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Data Field", "Classification", "Purpose", "Retained in Output"]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(10)

    fields = [
        ("Member Number", "PII — Indirect Identifier",
         "Loan-level tracking for migration analysis", "Audit worksheet only"),
        ("Current FICO Score", "Sensitive PII",
         "Current credit grade assignment", "Audit worksheet only"),
        ("Original FICO Score", "Sensitive PII",
         "Original credit grade for migration comparison", "Audit worksheet only"),
        ("Current Balance", "Sensitive Financial",
         "Reserve calculation basis", "Aggregated by pool"),
        ("Loan Pool", "Non-Sensitive",
         "Pool categorization for segmented analysis", "All report tabs"),
        ("Days Delinquent", "Sensitive Financial",
         "Delinquency rate calculation", "Aggregated by pool"),
        ("Open Date", "Low Sensitivity",
         "Life-of-loan window calculation", "Not output"),
        ("Interest Rate", "Sensitive Financial",
         "Weighted average rate analysis", "Not output"),
        ("Original Loan Amount", "Sensitive Financial",
         "Balance adjustment verification", "Not output"),
    ]
    for f in fields:
        add_table_row(table, f)

    doc.add_paragraph()
    add_heading_styled(doc, "3.2 Data NOT Collected", level=2)
    doc.add_paragraph(
        "The following categories of sensitive data are explicitly excluded from the "
        "system. They are never requested, ingested, stored, or processed:"
    )
    excluded = [
        "Social Security Numbers (SSN) or Tax Identification Numbers (TIN)",
        "Member names (first, last, or full)",
        "Dates of birth",
        "Physical addresses or mailing addresses",
        "Email addresses or phone numbers",
        "Employment information or income data",
        "Bank account numbers or routing numbers",
        "Full credit bureau reports or tradeline details",
        "Authentication credentials for member accounts",
    ]
    for item in excluded:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 4. DATA LIFECYCLE
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Data Lifecycle Management", level=1)

    add_heading_styled(doc, "4.1 Data Ingestion", level=2)
    doc.add_paragraph(
        "Source data files are provided by credit unions in Excel (.xlsx) or CSV format "
        "and stored on a secured network file share. The import process:"
    )
    ingestion = [
        "Reads only the specific columns defined in the client configuration — extraneous "
        "columns in source files are ignored and never loaded.",
        "Validates data types and ranges (e.g., FICO scores must be numeric, balances must "
        "be parseable as currency values).",
        "Strips account suffixes from member numbers to produce a standardized identifier.",
        "Assigns credit grades based on configured score ranges (e.g., 720+ = A+, "
        "660–719 = A).",
        "Calculates migration status by comparing current vs. original grades.",
    ]
    for item in ingestion:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, "4.2 Data Processing", level=2)
    doc.add_paragraph(
        "During report generation, loan-level data is processed through the CECL engine "
        "to produce regulatory metrics. The key processing steps include:"
    )
    processing = [
        "Risk Migration Matrix: Loans are cross-tabulated by original grade vs. current "
        "grade. Only aggregated dollar amounts and percentages are computed — no individual "
        "loan details persist.",
        "Reserve Calculations: Expected loss amounts are computed per loan using "
        "pool-specific reserve rates. Individual results are summed to pool-level totals "
        "for reporting.",
        "Historical Trend Analysis: Balance data from the WARM workbook is loaded at the "
        "pool and grade level — no member-level historical data is processed.",
        "Charge-Off and Recovery Analysis: Monthly totals are loaded from the WARM "
        "workbook's summary tabs — individual transaction details are not processed.",
    ]
    for item in processing:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, "4.3 Data at Rest", level=2)
    doc.add_paragraph(
        "Processed loan data is stored in a PostgreSQL database running on localhost. "
        "Key characteristics:"
    )
    at_rest = [
        "The database server binds to localhost (127.0.0.1) only — it is not accessible "
        "from any other machine on the network.",
        "Database credentials are stored in a local environment file (.env) that is not "
        "committed to version control.",
        "The database stores only the minimum fields required for CECL calculations "
        "(member number, FICO scores, balances, pool assignments, and calculated grades).",
        "Historical snapshots are retained for trend analysis across reporting periods.",
    ]
    for item in at_rest:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 5. SECURE DATA TRANSFER VIA EGNYTE
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Secure Data Transfer via Egnyte", level=1)
    doc.add_paragraph(
        "Credit union source data files are transferred to the analysis environment using "
        "Egnyte, an enterprise-grade cloud content management and file sharing platform. "
        "Founded in 2007 and headquartered in Mountain View, California, Egnyte serves "
        "over 16,000 business customers worldwide and is purpose-built for secure business "
        "file management, regulatory compliance, and controlled collaboration."
    )

    add_heading_styled(doc, "5.1 Egnyte Security Certifications & Compliance", level=2)
    doc.add_paragraph(
        "Egnyte maintains the following industry-recognized security certifications and "
        "compliance standards:"
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Certification / Standard", "Description"]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(10)

    certs = [
        ("SOC 2 Type II", "Independent audit verifying security, availability, and "
         "confidentiality controls over an extended period."),
        ("HIPAA Compliant", "Meets Health Insurance Portability and Accountability Act "
         "requirements for protected health information — demonstrating a standard of "
         "data protection that exceeds the requirements for financial data."),
        ("FINRA Compliant", "Meets Financial Industry Regulatory Authority requirements "
         "for financial records management and retention."),
        ("ISO/IEC 27001", "International standard for information security management "
         "systems (ISMS), covering risk assessment and security controls."),
        ("SSAE-16 (Tier II\nData Centers)", "Customer files are stored in Tier II "
         "SSAE-16 compliant data center facilities with physical security, redundancy, "
         "and environmental controls."),
    ]
    for cert in certs:
        add_table_row(table, cert)

    doc.add_paragraph()

    add_heading_styled(doc, "5.2 Encryption & Data Protection", level=2)
    encryption_items = [
        "Data at Rest: All files stored on Egnyte servers are protected with "
        "AES 256-bit encryption, the same standard used by financial institutions "
        "and government agencies.",
        "Data in Transit: All file transfers between client devices and Egnyte servers "
        "use TLS/SSL encryption, ensuring data cannot be intercepted during transmission.",
        "Enterprise Key Management (EKM): For enterprise-tier accounts, encryption keys "
        "can be managed externally via Amazon Cloud HSM or Microsoft Azure Key Vault, "
        "providing additional key custody control.",
    ]
    for item in encryption_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, "5.3 Access Controls & Authentication", level=2)
    access_items = [
        "Granular Permissions: File and folder access is controlled at a granular level — "
        "each credit union's data is isolated and accessible only to authorized personnel.",
        "Two-Factor Authentication (2FA): Egnyte supports mandatory two-factor "
        "authentication for all user accounts, adding a layer beyond password security.",
        "Single Sign-On (SSO): Integration with enterprise SSO providers ensures "
        "centralized identity management.",
        "Custom Password Policies: Administrators can enforce password strength, length, "
        "expiration, and reuse restrictions.",
        "Account Lockout: Automatic account lockout after repeated failed login attempts "
        "protects against brute-force attacks.",
        "Remote Device Wipe: In the event of device loss or compromise, data can be "
        "remotely wiped from computers and mobile devices.",
    ]
    for item in access_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, "5.4 Audit & Monitoring", level=2)
    audit_items = [
        "Comprehensive Audit Logging: Egnyte logs all activity including user logins, "
        "file uploads, downloads, shares, and permission changes.",
        "Sharing Controls: Shared links can be password-protected, set to expire "
        "automatically, and restricted to specific recipients.",
        "Abnormal Activity Detection: Enterprise accounts include suspicious login "
        "detection and alerts for unusual file access patterns.",
        "Ransomware Protection: Built-in detection against known ransomware threats "
        "with file versioning for recovery of affected files.",
    ]
    for item in audit_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, "5.5 Data Transfer Workflow", level=2)
    doc.add_paragraph(
        "The data transfer process between credit unions and the CECL analysis environment "
        "follows this secured workflow:"
    )
    workflow_steps = [
        "Credit union personnel upload source data files (AIRESLOANS exports, credit "
        "pull files, WARM workbooks) to their designated Egnyte folder.",
        "All file transfers are encrypted in transit via TLS/SSL and encrypted at rest "
        "via AES-256 upon arrival on Egnyte servers.",
        "Access to each credit union's folder is restricted to authorized personnel only, "
        "with no cross-client visibility.",
        "The CECL analyst downloads the source files from Egnyte to the secured local "
        "workstation for processing.",
        "Once downloaded, all data processing occurs entirely on the local workstation — "
        "no member data is transmitted back to Egnyte or any other external service.",
        "Completed reports may be uploaded back to Egnyte for secure delivery to the "
        "credit union, again protected by AES-256 encryption at rest and TLS/SSL in transit.",
    ]
    for i, step in enumerate(workflow_steps, 1):
        p = doc.add_paragraph(f"Step {i}: {step}")
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "At no point during the Egnyte transfer process is member data exposed to "
        "unauthorized parties. Egnyte's SOC 2 Type II certification provides independent "
        "assurance that these security controls are operating effectively."
    )
    run.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 6. ACCESS CONTROLS
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Access Controls & Authentication", level=1)

    add_heading_styled(doc, "6.1 System Access", level=2)
    controls = [
        "The application runs on a single workstation that requires Windows domain "
        "authentication to access.",
        "Source data files and output reports reside on a network share with "
        "role-based NTFS permissions restricting access to authorized personnel.",
        "The PostgreSQL database requires username/password authentication and is "
        "accessible only from the local machine. Database credentials are stored "
        "in Windows Credential Manager (encrypted via DPAPI) rather than in plaintext "
        "configuration files.",
        "No web-facing interfaces are exposed — the system has no public endpoints, "
        "APIs, or remote access capabilities.",
    ]
    for item in controls:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, "6.2 Principle of Least Privilege", level=2)
    doc.add_paragraph(
        "The system operates under the principle of least privilege:"
    )
    lp = [
        "Source files are read-only — the import process never modifies original data files.",
        "The database user has permissions limited to the cecl_migration_db database.",
        "Output reports are written to a designated Reports directory with controlled access.",
        "The application runs under a standard user account — no administrative "
        "privileges are required.",
    ]
    for item in lp:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 7. DATA PROCESSING SAFEGUARDS
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Data Processing Safeguards", level=1)

    add_heading_styled(doc, "7.1 Data Minimization", level=2)
    doc.add_paragraph(
        "The system follows data minimization principles throughout its pipeline:"
    )
    minimization = [
        "Only fields strictly necessary for CECL compliance calculations are extracted "
        "from source files. Additional columns present in source files are discarded "
        "during import.",
        "Member numbers are used solely as unique identifiers for tracking loan-level "
        "migration between periods — they are never used for member lookup or contact.",
        "FICO scores are immediately converted to categorical grades (A+, A, B, C, D) "
        "upon import. While raw scores are retained for audit purposes, all analytical "
        "processing uses the categorical grades.",
        "All report outputs present data at the aggregated pool level. The only exception "
        "is an audit detail worksheet described in Section 7.",
    ]
    for item in minimization:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, "7.2 No Automated Decision-Making", level=2)
    doc.add_paragraph(
        "The system does not make automated decisions about individual members. It produces "
        "aggregate statistical reports for regulatory compliance. No member is individually "
        "identified, scored, rated, or acted upon based on the system's output. All "
        "credit-related decisions remain with the credit union's authorized personnel."
    )

    add_heading_styled(doc, "7.3 Parameterized Database Queries", level=2)
    doc.add_paragraph(
        "All database queries use parameterized statements via SQLAlchemy's text() binding "
        "mechanism. No raw string concatenation is used in query construction, preventing "
        "SQL injection vulnerabilities."
    )

    add_heading_styled(doc, "7.4 Input Validation", level=2)
    doc.add_paragraph(
        "Source data undergoes validation during import:"
    )
    validation = [
        "Numeric fields (FICO scores, balances) are validated and coerced to appropriate types.",
        "Loan codes are mapped against a configured whitelist of valid pool assignments.",
        "Records with invalid or missing critical fields are flagged during import.",
        "Balance values are sanitized — currency symbols, commas, and parenthetical "
        "negatives are normalized before storage.",
    ]
    for item in validation:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 8. OUTPUT REPORT CONTROLS
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. Output Report Controls", level=1)

    add_heading_styled(doc, "8.1 Aggregated Report Tabs", level=2)
    doc.add_paragraph(
        "The vast majority of report content consists of aggregated, pool-level data "
        "that cannot be traced back to individual members:"
    )
    agg_tabs = [
        "Executive Summary — pool totals, reserve rates, portfolio-level metrics",
        "Risk Change by Credit Score — dollar and percentage migration matrices by grade",
        "Risk Change per Pool — same matrices segmented by loan pool",
        "Improved/Deteriorated Summary — pool-level counts and balances by migration status",
        "Environmental Factor by Pool — economic factor adjustments at pool level",
        "ACL Reserve by Pool — allowance calculations at pool level",
        "Historical Trend charts — pool-by-grade balance trends over time",
        "Charge-Off, Recovery, and Delinquency tabs — monthly pool-level totals",
        "Historical Balance Detail — grade-level balances per pool per month",
    ]
    for item in agg_tabs:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "None of the above tabs contain member numbers, individual FICO scores, or any "
        "data that could identify a specific member."
    )
    run.font.bold = True

    add_heading_styled(doc, "8.2 Audit Detail Worksheet", level=2)
    doc.add_paragraph(
        "Each report includes an \"All Loans\" worksheet that contains loan-level detail "
        "for audit verification purposes. This worksheet includes:"
    )
    audit_fields = [
        "Member Number (indirect identifier — not a name or SSN)",
        "Loan Pool assignment",
        "Current Balance",
        "Original and Current FICO Scores",
        "Original and Current Grade assignments",
        "Migration Status (Improved / Deteriorated / Unchanged)",
        "Reserve Rate and Expected Loss Amount",
    ]
    for item in audit_fields:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        "This worksheet exists solely to enable examiners and auditors to verify the "
        "accuracy of the aggregated calculations. It allows tracing any pool-level total "
        "back to the individual loans that comprise it. The member number serves as a "
        "cross-reference key — it does not reveal the member's identity without access "
        "to the credit union's separate core system."
    )
    doc.add_paragraph(
        "This worksheet is automatically password-protected when the report is generated, "
        "preventing casual modification or accidental exposure of member-level data. "
        "The protection password is provided separately to authorized examiners."
    )

    add_heading_styled(doc, "8.3 Report Distribution", level=2)
    doc.add_paragraph(
        "Generated reports are saved to a secured network share directory (Reports/). "
        "Distribution of reports is handled through the credit union's existing secure "
        "document delivery processes. Reports are not emailed, uploaded to cloud storage, "
        "or transmitted via any automated mechanism."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 9. EXTERNAL DATA TRANSMISSION
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. External Data Transmission", level=1)
    doc.add_paragraph(
        "The system makes limited outbound HTTP requests to retrieve publicly available "
        "economic data for environmental factor calculations. These requests are strictly "
        "one-directional and carry zero member data."
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Data Source", "Endpoint", "Data Sent", "PII Transmitted"]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(10)

    ext_sources = [
        ("Bureau of Labor\nStatistics (BLS)", "api.bls.gov\n/publicAPI/v1/", "State FIPS code,\nBLS series ID", "None"),
        ("U.S. Census Bureau", "api.census.gov\n/data/{year}/acs/", "State FIPS code", "None"),
        ("U.S. Courts\n(Bankruptcy Stats)", "uscourts.gov\n/sites/default/files/", "No parameters\n(public file download)", "None"),
    ]
    for src in ext_sources:
        add_table_row(table, src)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "No member numbers, FICO scores, balances, or any other member-specific data "
        "is included in any outbound request. The system only downloads aggregate "
        "economic statistics (unemployment rates, median income, bankruptcy filing counts) "
        "for use as environmental adjustment factors."
    )
    run.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 10. RETENTION & DISPOSAL
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. Data Retention & Disposal", level=1)

    add_heading_styled(doc, "10.1 Database Records", level=2)
    doc.add_paragraph(
        "Loan snapshots are retained in the local database to support quarter-over-quarter "
        "trend analysis required by CECL methodology. Historical data enables the calculation "
        "of lifetime loss rates and migration trend charts. Records are retained in "
        "accordance with the credit union's data retention policy."
    )

    add_heading_styled(doc, "10.2 Source Files", level=2)
    doc.add_paragraph(
        "Original source files (AIRESLOANS exports, credit pull files, WARM workbooks) "
        "are maintained on the secured network share in client-specific, date-organized "
        "folders. These files are subject to the credit union's document retention schedule."
    )

    add_heading_styled(doc, "10.3 Generated Reports", level=2)
    doc.add_paragraph(
        "Output reports are retained in the Reports directory. Prior versions may be "
        "moved to an Archive folder when superseded. Report retention follows the credit "
        "union's regulatory document retention requirements."
    )

    add_heading_styled(doc, "10.4 Temporary Files", level=2)
    doc.add_paragraph(
        "The main report generation pipeline does not create persistent temporary files. "
        "All intermediate calculations are performed in-memory using pandas DataFrames "
        "that are released when processing completes. No member data is written to "
        "system temp directories during normal operation."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 11. IMPLEMENTED CONTROLS & REMAINING RECOMMENDATIONS
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "11. Security Controls & Recommendations", level=1)

    # ── Implemented ──
    add_heading_styled(doc, "11.1 Implemented Controls", level=2)
    doc.add_paragraph(
        "The following security enhancements have been implemented in the current system:"
    )

    add_heading_styled(doc, "11.1.1 Credential Management — Windows Credential Manager", level=3)
    doc.add_paragraph(
        "Database credentials have been migrated from the plaintext .env file to "
        "Windows Credential Manager (via the keyring library). The system retrieves "
        "the DATABASE_URL securely at runtime through the Windows Data Protection API "
        "(DPAPI), which encrypts credentials using the logged-in user's Windows account. "
        "The .env file is retained as a fallback for environments where keyring is "
        "unavailable, but the Credential Manager is checked first."
    )
    doc.add_paragraph(
        "Management utility: cecl_credentials.py (--store, --verify, --delete)"
    )

    add_heading_styled(doc, "11.1.2 Audit Detail Worksheet Protection", level=3)
    doc.add_paragraph(
        "The \"All Loans\" audit detail worksheet is now automatically password-protected "
        "when reports are generated. Sheet protection prevents casual modification or "
        "accidental exposure of member-level data (member numbers, FICO scores, and "
        "individual balances). Authorized examiners can access the data using the "
        "protection password provided separately."
    )

    add_heading_styled(doc, "11.1.3 Access Logging", level=3)
    doc.add_paragraph(
        "A centralized audit logging system has been implemented (cecl_audit_log.py). "
        "All report generation and data import operations are logged to a rotating log "
        "file (logs/cecl_audit.log) with the following fields:"
    )
    log_fields = [
        "Timestamp — date and time of the operation",
        "User — Windows username of the operator (captured via getpass)",
        "Operation — type of action (REPORT_GENERATED, DATA_IMPORTED, SESSION_START/END)",
        "Client — which credit union configuration was used",
        "Details — snapshot date, report type, output filename, record counts, success/failure",
    ]
    for item in log_fields:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph(
        "Log files use a rotating file handler: 5 MB maximum per file with 10 backups "
        "retained (50 MB total). This provides a verifiable access trail for regulatory "
        "review without unbounded storage growth."
    )

    add_heading_styled(doc, "11.1.4 Data Retention Automation", level=3)
    doc.add_paragraph(
        "An automated data retention tool has been implemented (cecl_retention.py) that "
        "can purge database records and report files beyond a configurable retention "
        "period (default: 7 years / 84 months). The tool operates in two modes:"
    )
    retention_items = [
        "Dry Run (--dry-run): Previews what would be deleted without making changes.",
        "Execute (--execute): Deletes old records after explicit user confirmation, "
        "with all deletions logged to the audit log.",
    ]
    for item in retention_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph(
        "The tool can independently purge database records, archived source files, and "
        "old report files. This reduces the volume of stored PII and limits exposure in "
        "the event of a breach."
    )

    # ── Remaining ──
    add_heading_styled(doc, "11.2 Remaining Recommendations", level=2)
    doc.add_paragraph(
        "The following items are recommended for future implementation:"
    )

    add_heading_styled(doc, "11.2.1 Database Encryption at Rest", level=3)
    doc.add_paragraph(
        "Currently, the PostgreSQL database stores data unencrypted on disk. Enabling "
        "full-disk encryption (e.g., Windows BitLocker) on the workstation would protect "
        "data at rest against physical theft or unauthorized disk access. This is an "
        "operating system-level configuration that should be coordinated with IT."
    )

    add_heading_styled(doc, "11.2.2 Member Number Hashing", level=3)
    doc.add_paragraph(
        "For environments with heightened security requirements, member numbers stored "
        "in the database could be replaced with salted hashes. The original member numbers "
        "would only appear during the import-to-report pipeline and not be persisted. "
        "This would be a significant architectural change and should be evaluated against "
        "the practical need for cross-period loan tracking."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # APPENDIX A
    # ═══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Appendix A: Data Field Inventory", level=1)
    doc.add_paragraph(
        "Complete inventory of data fields processed by the system, their classification, "
        "and their disposition in output reports."
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Field", "Source", "Classification", "Stored in DB", "In Output Report"]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(9)

    inventory = [
        ("member_number", "AIRES export", "PII (indirect)", "Yes", "Audit sheet only"),
        ("current_fico_score", "Credit pull / WARM", "Sensitive PII", "Yes", "Audit sheet only"),
        ("original_fico_score", "Credit pull / WARM", "Sensitive PII", "Yes", "Audit sheet only"),
        ("current_balance", "AIRES export", "Sensitive Financial", "Yes", "Aggregated by pool"),
        ("original_loan_amount", "AIRES export", "Sensitive Financial", "Yes", "Not in output"),
        ("interest_rate", "AIRES export", "Sensitive Financial", "Yes", "Not in output"),
        ("open_date", "AIRES export", "Low Sensitivity", "Yes", "Not in output"),
        ("days_delinquent", "AIRES export", "Sensitive Financial", "Yes", "Aggregated by pool"),
        ("loan_pool", "Derived from code", "Non-Sensitive", "Yes", "All tabs"),
        ("current_grade", "Derived from FICO", "Derived PII", "Yes", "Aggregated by grade"),
        ("original_grade", "Derived from FICO", "Derived PII", "Yes", "Aggregated by grade"),
        ("migration_status", "Calculated", "Derived", "Yes", "Aggregated by pool"),
        ("reserve_rate", "Config + engine", "Non-Sensitive", "Yes", "Pool-level tabs"),
        ("expected_loss_amount", "Calculated", "Derived", "Yes", "Aggregated by pool"),
        ("credit_union", "Config", "Non-Sensitive", "Yes", "Report headers"),
        ("snapshot_date", "Config", "Non-Sensitive", "Yes", "Report headers"),
    ]
    for inv in inventory:
        add_table_row(table, inv)

    # ── Footer on all pages ──
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CECL Credit Migration Analysis — Data Security & Privacy Policy — CONFIDENTIAL")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


if __name__ == '__main__':
    out_path = os.path.join(os.path.dirname(__file__), 'Reports',
                            'CECL_Data_Security_Privacy_Policy.docx')
    doc = build_document()
    doc.save(out_path)
    print(f"Saved: {out_path}")
